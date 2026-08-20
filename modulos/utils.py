import pandas as pd
import re
import os
import smtplib
import io
import base64
import requests
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import streamlit as st
import datetime

def converter_para_numero(valor):
    if isinstance(valor, (int, float)): return float(valor)
    texto = str(valor).replace("R$", "").replace("%", "").strip()
    if texto == "": return 0.0
    if "," in texto:
        if "." in texto: texto = texto.replace(".", "")
        texto = texto.replace(",", ".")
    try: return float(texto)
    except: return 0.0

def padronizar_nome(texto):
    if not texto: return ""
    excecoes = ['de', 'da', 'do', 'das', 'dos', 'e']
    palavras = str(texto).strip().split()
    return " ".join([p.lower() if p.lower() in excecoes else p.lower().capitalize() for p in palavras])

def padronizar_telefone(tel):
    if not tel: return ""
    num = re.sub(r'\D', '', str(tel))
    if len(num) == 11: return f"({num[:2]}) {num[2]} {num[3:7]}-{num[7:]}"
    elif len(num) == 10: return f"({num[:2]}) {num[2:6]}-{num[6:]}"
    return tel 

def extrair_tabela_crm_itens(itens_str):
    dados_tabela = []
    if not itens_str or itens_str == 'nan': return dados_tabela
    for elem in str(itens_str).split(";"):
        elem = elem.strip()
        if not elem: continue
        cod, qtd, nome, v_u = "-", 1, elem, 0.0
        if "[Cód:" in elem:
            try:
                parte_qtd_nome, rest = elem.split("[Cód:", 1)
                cod, rest2 = rest.split("]", 1)[0].strip(), rest.split("]", 1)[1]
                if "x " in parte_qtd_nome:
                    qtd = int(parte_qtd_nome.split("x ", 1)[0].strip())
                    nome = parte_qtd_nome.split("x ", 1)[1].strip()
                else: nome = parte_qtd_nome.strip()
                if "(R$" in rest2: v_u = converter_para_numero(rest2.split("(R$", 1)[1].replace(")", "").strip())
            except: pass
        elif "x " in elem:
            try:
                qtd = int(elem.split("x ", 1)[0].strip())
                nome = elem.split("x ", 1)[1].strip()
            except: pass
        subtotal = v_u * qtd
        dados_tabela.append({"Código KME": cod, "Produto / Serviço": nome, "Qtd": qtd, "Valor Unit.": f"R$ {v_u:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".") if v_u > 0 else "-", "Subtotal": f"R$ {subtotal:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".") if subtotal > 0 else "-"})
    return dados_tabela

def validar_inconsistencias_carrinho(carrinho, df_regras):
    avisos = []
    if df_regras.empty or not carrinho: return avisos
    codigos_no_carrinho = [str(item.get('codigo', '')).strip().lstrip('0') for item in carrinho]
    for _, regra in df_regras.iterrows():
        gatilho = str(regra.get('Item_Gatilho', '')).strip().lstrip('0')
        exigidos_str = str(regra.get('Itens_Exigidos', '')).strip()
        msg = str(regra.get('Mensagem_Aviso', 'Inconsistência detectada.'))
        tipo_regra = str(regra.get('Tipo_Regra', 'Exigencia')).strip().lower()
        if gatilho in codigos_no_carrinho:
            itens_relacionados = [c.strip().lstrip('0') for c in exigidos_str.split(';') if c.strip()]
            if 'exig' in tipo_regra:
                if not any(codigo in codigos_no_carrinho for codigo in itens_relacionados): avisos.append(msg)
            elif 'incompat' in tipo_regra or 'proib' in tipo_regra:
                if any(codigo in codigos_no_carrinho for codigo in itens_relacionados): avisos.append(msg)
    return avisos

def obter_detalhes_split(row_data, df_produtos, df_valor_sensor, df_valor_ponto, unidade_selecionada):
    itens_str = str(row_data.get('Itens_Orcamento', ''))
    desc_p, desc_a, desc_i = converter_para_numero(row_data.get('Desc_Prod', '0')) / 100, converter_para_numero(row_data.get('Desc_Alarme', '0')) / 100, converter_para_numero(row_data.get('Desc_Imagem', '0')) / 100
    bruto_prod, bruto_alarme, bruto_imagem, mao_obra = 0.0, 0.0, 0.0, 0.0
    itens_parsed, qtd_abertura, qtd_ivp = [], 0, 0
    
    for item in itens_str.split(";"):
        if "x " in item:
            try:
                qtd = int(item.strip().split("x ", 1)[0])
                nome_item = item.strip().split("x ", 1)[1].split("[Cód:")[0].strip() if "[Cód:" in item else item.strip().split("x ", 1)[1].strip()
            except: qtd, nome_item = 0, ""
            prod_info = df_produtos[df_produtos['Nome_Item'].astype(str).str.strip() == nome_item]
            if not prod_info.empty:
                prod = prod_info.iloc[0]
                ts = str(prod.get('Tipo_Sensor', '')).strip().upper()
                if ts == 'ABERTURA': qtd_abertura += qtd
                elif ts == 'IVP': qtd_ivp += qtd
                itens_parsed.append({'qtd': qtd, 'prod': prod})
                
    for it in itens_parsed:
        qtd, prod = it['qtd'], it['prod']
        cat, grupo, cod_item = str(prod.get('Categoria_Receita', '')).strip().lower(), str(prod.get('Grupo_Itens', '')).strip().lower(), str(prod.get('Codigo_KME', '')).strip().lstrip('0')
        v_u = converter_para_numero(prod.get('Preco_Venda', 0)) if converter_para_numero(prod.get('Preco_Venda', 0)) > 0 else converter_para_numero(prod.get('Preco_LOC_36', 0))
        
        if ("obra" in cat or "instala" in cat) and not df_valor_ponto.empty:
            match_mo = df_valor_ponto[(df_valor_ponto['Unidade'].astype(str).str.strip() == unidade_selecionada) & (df_valor_ponto['Nome_Item'].astype(str).str.strip() == str(prod.get('Nome_Item', '')).strip())]
            if not match_mo.empty: v_u = converter_para_numero(match_mo.iloc[0]['Valor_MO'])
            
        if cod_item in ['254000000042', '254000000377', '25400000042', '25400000377']:
            if not df_valor_sensor.empty:
                match = df_valor_sensor[(df_valor_sensor['Codigo_Servico'].astype(str).str.strip().str.lstrip('0') == cod_item) & (pd.to_numeric(df_valor_sensor['Sensor_Abertura'], errors='coerce') == qtd_abertura) & (pd.to_numeric(df_valor_sensor['Sensor_IVP'], errors='coerce') == qtd_ivp)]
                if not match.empty: v_u = converter_para_numero(match.iloc[0]['Preco'])
                    
        if "obra" in cat or "instala" in cat: mao_obra += (v_u * qtd)
        elif "produto" in cat or "equipamento" in cat: bruto_prod += (v_u * qtd)
        else:
            if "imagem" in grupo: bruto_imagem += (v_u * qtd)
            else: bruto_alarme += (v_u * qtd)
            
    liq_prod = bruto_prod * (1 - desc_p)
    liq_mrr = (bruto_alarme * (1 - desc_a)) + (bruto_imagem * (1 - desc_i))

    mrr_fmt = f"R$ {liq_mrr:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    eqp_fmt = f"R$ {liq_prod:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    mo_fmt = f"R$ {mao_obra:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    
    return mrr_fmt, eqp_fmt, mo_fmt

def calcular_novos_valores_proposta(row_data, df_produtos, df_valor_sensor):
    itens_str = str(row_data.get('Itens_Orcamento', ''))
    desc_p, desc_a, desc_i = converter_para_numero(row_data.get('Desc_Prod', '0')) / 100, converter_para_numero(row_data.get('Desc_Alarme', '0')) / 100, converter_para_numero(row_data.get('Desc_Imagem', '0')) / 100
    bruto_prod, bruto_alarme, bruto_imagem, mao_obra = 0.0, 0.0, 0.0, 0.0
    itens_parsed, qtd_abertura, qtd_ivp = [], 0, 0
    
    for item in itens_str.split(";"):
        if "x " in item:
            try:
                qtd = int(item.strip().split("x ", 1)[0])
                nome_item = item.strip().split("x ", 1)[1].split("[Cód:")[0].strip() if "[Cód:" in item else item.strip().split("x ", 1)[1].strip()
            except: qtd, nome_item = 0, ""
            prod_info = df_produtos[df_produtos['Nome_Item'].astype(str).str.strip() == nome_item]
            if not prod_info.empty:
                prod = prod_info.iloc[0]
                ts = str(prod.get('Tipo_Sensor', '')).strip().upper()
                if ts == 'ABERTURA': qtd_abertura += qtd
                elif ts == 'IVP': qtd_ivp += qtd
                itens_parsed.append({'qtd': qtd, 'prod': prod})
                
    for it in itens_parsed:
        qtd, prod = it['qtd'], it['prod']
        cat, grupo, cod_item = str(prod.get('Categoria_Receita', '')).strip().lower(), str(prod.get('Grupo_Itens', '')).strip().lower(), str(prod.get('Codigo_KME', '')).strip().lstrip('0')
        v_u = converter_para_numero(prod.get('Preco_Venda', 0)) if converter_para_numero(prod.get('Preco_Venda', 0)) > 0 else converter_para_numero(prod.get('Preco_LOC_36', 0))
        
        if cod_item in ['254000000042', '254000000377', '25400000042', '25400000377']:
            if not df_valor_sensor.empty:
                match = df_valor_sensor[(df_valor_sensor['Codigo_Servico'].astype(str).str.strip().str.lstrip('0') == cod_item) & (pd.to_numeric(df_valor_sensor['Sensor_Abertura'], errors='coerce') == qtd_abertura) & (pd.to_numeric(df_valor_sensor['Sensor_IVP'], errors='coerce') == qtd_ivp)]
                if not match.empty: v_u = converter_para_numero(match.iloc[0]['Preco'])
                    
        if "obra" in cat or "instala" in cat: mao_obra += (v_u * qtd)
        elif "produto" in cat or "equipamento" in cat: bruto_prod += (v_u * qtd)
        else:
            if "imagem" in grupo: bruto_imagem += (v_u * qtd)
            else: bruto_alarme += (v_u * item['quantidade'])
            
    novo_total_mrr, novo_total_setup = (bruto_alarme * (1 - desc_a)) + (bruto_imagem * (1 - desc_i)), (bruto_prod * (1 - desc_p)) + mao_obra
    return f"R$ {novo_total_mrr:,.2f}".replace(",", "_").replace(".", ",").replace("_", "."), f"R$ {novo_total_setup:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")

def obter_emails_gestores(df_users, unidade_user, vertical_user):
    emails = []
    lideres = df_users[(df_users['Perfil'].astype(str).str.strip() == 'Lider') & (df_users['Unidade'].astype(str).str.strip().str.lower() == str(unidade_user).strip().lower())]
    emails.extend(lideres['Email_C'].tolist())
    
    if "varejo" in str(vertical_user).lower():
        gerentes = df_users[df_users['Perfil'].astype(str).str.strip() == 'Gerente_Varejo']
        emails.extend(gerentes['Email_C'].tolist())
    elif "condominio" in str(vertical_user).lower():
        gerentes = df_users[df_users['Perfil'].astype(str).str.strip() == 'Gerente_Condominio']
        emails.extend(gerentes['Email_C'].tolist())
        
    return list(set(emails))

def enviar_email_aprovacao(nome_consultor, unidade, vertical, valor_mrr, valor_equip, valor_mo, emails_destino):
    try:
        if "smtp" not in st.secrets:
            st.info("💡 E-mail de aprovação gerado! (Configure SMTP no Cloud para envio real).")
            return True
            
        remetente = st.secrets["smtp"]["email"]
        senha = st.secrets["smtp"]["password"]
        servidor = st.secrets["smtp"]["server"]
        porta = st.secrets["smtp"]["port"]

        msg = MIMEMultipart()
        msg['From'] = remetente
        msg['To'] = ", ".join(emails_destino)
        msg['Subject'] = "Contrato Aprovado pelo Cliente 🏆"

        html = f"""
        <div style="font-family: Arial, sans-serif; color: #1e293b;">
            <h2 style="color: #0066cc;">Nova Aprovação de Contrato!</h2>
            <p>Olá, Sinalizamos a aprovação do contrato abaixo:</p>
            <table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse; width: 100%; max-width: 600px;">
                <tr style="background-color: #f8fafc;"><td style="width: 40%;"><b>Nome Consultor:</b></td><td>{nome_consultor}</td></tr>
                <tr><td><b>Unidade:</b></td><td>{unidade}</td></tr>
                <tr style="background-color: #f8fafc;"><td><b>Segmento:</b></td><td>{vertical}</td></tr>
                <tr><td><b>Valor Mensalidade:</b></td><td><span style="color: #059669; font-weight: bold;">{valor_mrr}</span></td></tr>
                <tr style="background-color: #f8fafc;"><td><b>Valor Venda Equipamentos:</b></td><td>{valor_equip}</td></tr>
                <tr><td><b>Valor Mão de Obra:</b></td><td>{valor_mo}</td></tr>
            </table>
            <br><p><i>Obs: Proposta Segue para Análise de Cadastro/Crédito.</i></p>
        </div>
        """
        msg.attach(MIMEText(html, 'html'))
        server = smtplib.SMTP(servidor, porta)
        server.starttls()
        server.login(remetente, senha)
        server.sendmail(remetente, emails_destino, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"Erro ao disparar email: {e}")
        return False

def enviar_email_proposta_cliente(nome_cliente, email_destino, html_conteudo):
    try:
        if "smtp" not in st.secrets:
            st.info("💡 E-mail para o cliente gerado com sucesso! (Configure SMTP no Cloud para envio real).")
            return True
            
        remetente = st.secrets["smtp"]["email"]
        senha = st.secrets["smtp"]["password"]
        servidor = st.secrets["smtp"]["server"]
        porta = st.secrets["smtp"]["port"]

        msg = MIMEMultipart()
        msg['From'] = remetente
        msg['To'] = email_destino
        msg['Subject'] = f"Khronos - Sua Proposta Comercial"

        msg.attach(MIMEText(html_conteudo, 'html'))

        server = smtplib.SMTP(servidor, porta)
        server.starttls()
        server.login(remetente, senha)
        server.sendmail(remetente, [email_destino], msg.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"Erro ao disparar email para o cliente: {e}")
        return False

def enviar_email_recuperacao_senha(email_destino, senha_atual):
    try:
        if "smtp" not in st.secrets:
            # Removida a exibição da senha na tela por motivos de segurança!
            st.info("💡 E-mail de recuperação gerado com sucesso! (Configure SMTP no Cloud para envio real).")
            return True
            
        remetente = st.secrets["smtp"]["email"]
        senha = st.secrets["smtp"]["password"]
        servidor = st.secrets["smtp"]["server"]
        porta = st.secrets["smtp"]["port"]

        msg = MIMEMultipart()
        msg['From'] = remetente
        msg['To'] = email_destino
        msg['Subject'] = "Khronos Sales - Recuperação de Senha"

        html = f"""
        <div style="font-family: Arial, sans-serif; color: #1e293b;">
            <h2 style="color: #e20613;">Recuperação de Senha</h2>
            <p>Olá,</p>
            <p>Você solicitou a recuperação da sua senha de acesso ao portal Khronos Sales.</p>
            <p>Sua senha atual é: <strong>{senha_atual}</strong></p>
            <br>
            <p><i>Recomendamos que você anote sua senha em um local seguro.</i></p>
        </div>
        """
        msg.attach(MIMEText(html, 'html'))
        server = smtplib.SMTP(servidor, porta)
        server.starttls()
        server.login(remetente, senha)
        server.sendmail(remetente, [email_destino], msg.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"Erro ao disparar email de recuperação: {e}")
        return False

def gerar_html_proposta(cliente, proposta, vendedor, itens_carrinho, mrr, setup, condicao_texto):
    hoje = datetime.datetime.now().strftime("%d/%m/%Y")
    linhas_tabela = ""
    for item in itens_carrinho:
        qtd = item.get('quantidade', item.get('Qtd', 1))
        nome = item.get('nome', item.get('Produto / Serviço', ''))
        linhas_tabela += f"<tr><td style='text-align:center;'>{qtd}</td><td>{nome}</td></tr>"

    html = f"""
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <title>Proposta Comercial - {cliente}</title>
        <style>
            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: #1e293b; max-width: 800px; margin: 0 auto; padding: 40px; line-height: 1.6; }}
            .cabecalho {{ border-bottom: 3px solid #e20613; padding-bottom: 20px; margin-bottom: 30px; display: flex; justify-content: space-between; align-items: center; }}
            .logo-text {{ font-size: 28px; font-weight: 900; color: #e20613; margin: 0; }}
            .info-geral p {{ margin: 5px 0; font-size: 14px; color: #64748b; }}
            .box-cliente {{ background-color: #f8fafc; border-left: 4px solid #0f172a; padding: 20px; border-radius: 8px; margin-bottom: 30px; }}
            .box-cliente h2 {{ margin-top: 0; color: #0f172a; font-size: 20px; }}
            .box-cliente p {{ margin: 5px 0; font-size: 16px; }}
            table {{ width: 100%; border-collapse: collapse; margin-bottom: 30px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
            th, td {{ padding: 12px 15px; border-bottom: 1px solid #e2e8f0; text-align: left; }}
            th {{ background-color: #f1f5f9; color: #334155; font-weight: 600; text-transform: uppercase; font-size: 13px; }}
            .box-financeiro {{ background-color: #f0fdf4; border: 1px solid #bbf7d0; padding: 25px; border-radius: 8px; margin-bottom: 40px; }}
            .box-financeiro h3 {{ margin-top: 0; color: #166534; font-size: 18px; margin-bottom: 15px; border-bottom: 1px solid #bbf7d0; padding-bottom: 10px; }}
            .linha-valor {{ display: flex; justify-content: space-between; font-size: 16px; margin-bottom: 10px; }}
            .linha-destaque {{ display: flex; justify-content: space-between; font-size: 20px; font-weight: bold; color: #166534; margin-top: 15px; padding-top: 15px; border-top: 2px dashed #bbf7d0; }}
            .rodape {{ text-align: center; font-size: 12px; color: #94a3b8; border-top: 1px solid #e2e8f0; padding-top: 20px; }}
        </style>
    </head>
    <body>
        <div class="cabecalho">
            <div>
                <h1 class="logo-text">Khronos</h1>
                <p style="margin: 0; color: #64748b; font-weight: 600;">Proposta Comercial</p>
            </div>
            <div class="info-geral" style="text-align: right;">
                <p><b>Data:</b> {hoje}</p>
                <p><b>Consultor:</b> {vendedor}</p>
            </div>
        </div>
        
        <div class="box-cliente">
            <h2>Dados do Cliente</h2>
            <p><b>Nome/Razão Social:</b> {cliente}</p>
            <p><b>Referência do Projeto:</b> {proposta}</p>
        </div>
        
        <h3 style="color: #334155; margin-bottom: 15px;">Escopo do Projeto</h3>
        <table>
            <thead>
                <tr>
                    <th style="width: 10%; text-align: center;">Qtd</th>
                    <th>Produto / Serviço</th>
                </tr>
            </thead>
            <tbody>
                {linhas_tabela}
            </tbody>
        </table>
        
        <div class="box-financeiro">
            <h3>Resumo de Investimento</h3>
            <div class="linha-valor">
                <span>Investimento Inicial (Setup / Equipamentos):</span>
                <strong>{setup}</strong>
            </div>
            <div class="linha-valor" style="font-size: 14px; color: #475569;">
                <span>Condição de Pagamento:</span>
                <span>{condicao_texto}</span>
            </div>
            <div class="linha-destaque">
                <span>Total em Serviços (Mensalidade):</span>
                <span>{mrr} / mês</span>
            </div>
        </div>
        
        <div class="rodape">
            <p>Proposta comercial válida por 10 dias a partir da data de emissão.</p>
            <p><i>* Sujeita à aprovação de crédito e viabilidade técnica no local.</i></p>
        </div>
    </body>
    </html>
    """
    return html

def gerar_documento_contrato(lead_dados, mrr_formatado, setup_formatado, condicao_txt):
    try:
        from docx import Document
    except ImportError:
        return None, "Biblioteca 'python-docx' não instalada."

    caminho_template = "template_contrato.docx"
    if not os.path.exists(caminho_template):
        return None, "Arquivo 'template_contrato.docx' não encontrado na raiz."

    try:
        doc = Document(caminho_template)
    except Exception as e:
        return None, f"Erro ao abrir o arquivo Word original: {e}"

    hoje = datetime.datetime.now().strftime("%d/%m/%Y")
    
    substituicoes = {
        "{{NOME_CLIENTE}}": str(lead_dados.get("nome", "")),
        "{{CPF_CNPJ}}": str(lead_dados.get("cpf_cnpj", "")),
        "{{ENDERECO}}": f"{lead_dados.get('endereco', '')}, {lead_dados.get('numero', '')}",
        "{{CIDADE}}": str(lead_dados.get("cidade", "")),
        "{{ESTADO}}": str(lead_dados.get("estado", "")),
        "{{TELEFONE}}": str(lead_dados.get("telefone", "")),
        "{{EMAIL}}": str(lead_dados.get("email_cliente", "")),
        "{{VALOR_MENSAL}}": str(mrr_formatado),
        "{{VALOR_SETUP}}": str(setup_formatado),
        "{{CONDICAO_PGTO}}": str(condicao_txt),
        "{{DATA_ATUAL}}": hoje
    }

    def substituir_nas_runs(paragraphs):
        for p in paragraphs:
            for tag, valor in substituicoes.items():
                if tag in p.text:
                    for run in p.runs:
                        if tag in run.text:
                            run.text = run.text.replace(tag, valor)
                    if tag in p.text:
                        p.text = p.text.replace(tag, valor)

    substituir_nas_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_nas_runs(cell.paragraphs)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue(), None

def enviar_para_zapsign(docx_bytes, nome_cliente, email_cliente, telefone_cliente):
    try:
        if "zapsign" not in st.secrets or "token" not in st.secrets["zapsign"]:
            return False, "⚠️ Token da ZapSign não configurado no painel do Streamlit."

        token = st.secrets["zapsign"]["token"]
        url = f"https://api.zapsign.com.br/api/v1/docs/?api_token={token}"

        base64_doc = base64.b64encode(docx_bytes).decode('utf-8')
        
        tel_formatado = re.sub(r'\D', '', str(telefone_cliente))
        
        signer = {"name": nome_cliente}
        
        if email_cliente:
            signer["email"] = email_cliente
            signer["send_via"] = "email"
            
        if len(tel_formatado) >= 10:
            signer["phone_country"] = "55"
            signer["phone_number"] = tel_formatado
            if not email_cliente:
                signer["send_via"] = "whatsapp"

        payload = {
            "name": f"Contrato Khronos - {nome_cliente}",
            "base64_pdf": f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64_doc}",
            "signers": [signer]
        }

        response = requests.post(url, json=payload, headers={"Content-Type": "application/json"})

        if response.status_code == 200:
            return True, "Enviado com Sucesso para o cliente assinar!"
        else:
            return False, f"Falha ZapSign: Verifique seu Token."
            
    except Exception as e:
        return False, f"Erro na integração: {str(e)}"
